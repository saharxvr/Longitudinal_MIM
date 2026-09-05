"""Build v2 of the observer-variability paper (.docx): HUMAN READERS ONLY (no model).

Differences from v1:
  - No model anywhere (figures, text, discussion, supplementary).
  - Human-only 5x5 PAI heatmaps and human-only leave-one-out figure.
    - New Section 3.5 + Figure + Table: reader-subgroup coverage analysis.
    - Conclusions distinguish the observed coverage trade-off from a formal
        threshold for an adequate reference standard.
  - [PLACEHOLDERS] are highlighted by category with a color legend:
        yellow  = data to add before submission
        green   = optional suggestion / analysis to consider
        cyan    = editorial decision to confirm
    Author placeholders as [CAT|text], e.g. [DATA|N patients].

Run (from repo root):
    python python_files/Sahar_work/build_ov_paper_docx_v2.py
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX


REPO = Path(__file__).resolve().parents[1]
HUM = REPO / "Sahar_work" / "files" / "ov_results_human_only"
OUT_DOCX = REPO / "Sahar_work" / "OV_paper_draft_JThoracImaging_v3.docx"

HL = {
    "DATA": WD_COLOR_INDEX.YELLOW,
    "SUGGEST": WD_COLOR_INDEX.BRIGHT_GREEN,
    "DECIDE": WD_COLOR_INDEX.TURQUOISE,
}
TAG_RE = re.compile(r"\[(DATA|SUGGEST|DECIDE)\|(.*?)\]")


def _add_runs(p, text: str) -> None:
    """Emit runs for `text`, highlighting [CAT|...] placeholders by category."""
    pos = 0
    for m in TAG_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        cat, inner = m.group(1), m.group(2)
        r = p.add_run(f"[{inner}]")
        r.font.highlight_color = HL[cat]
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def body(doc, text: str) -> None:
    _add_runs(doc.add_paragraph(), text)


def labeled(doc, label: str, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label} ")
    r.bold = True
    _add_runs(p, text)


def bullets(doc, items: list[str]) -> None:
    for it in items:
        _add_runs(doc.add_paragraph(style="List Bullet"), it)


def title(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)


def center(doc, text: str, italic=False, size=11) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(p, text)
    for r in p.runs:
        r.italic = italic
        r.font.size = Pt(size)


def table(doc, headers, rows) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        c.paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)


def figure(doc, path: Path, caption: str, width_in: float) -> None:
    if not path.exists():
        doc.add_paragraph(f"[MISSING FIGURE: {path.name}]")
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


def legend_box(doc) -> None:
    p = doc.add_paragraph()
    p.add_run("Placeholder color key: ").bold = True
    r1 = p.add_run(" TO ADD — real data/values required before submission ")
    r1.font.highlight_color = HL["DATA"]
    p.add_run("   ")
    r2 = p.add_run(" SUGGESTION — optional analysis/addition to consider ")
    r2.font.highlight_color = HL["SUGGEST"]
    p.add_run("   ")
    r3 = p.add_run(" DECISION — confirm wording / choice ")
    r3.font.highlight_color = HL["DECIDE"]


def build() -> None:
    metrics = json.loads((HUM / "human_ov_metrics.json").read_text(encoding="utf-8"))
    cov = metrics["subgroup_coverage"]
    mg = metrics["subgroup_marginal_gain"]["all"]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)

    title(doc, "How Many Radiologists Define the Truth? Inter-observer Variability in the "
               "Detection of Interval Change on Longitudinal Chest Radiographs")
    center(doc, "[DATA|Author list]", size=11)
    center(doc, "[DATA|School of Computer Science and Engineering, The Hebrew University of Jerusalem; "
                "Department of Radiology, Hadassah Hebrew University Medical Center, Jerusalem, Israel]",
           italic=True, size=9)
    center(doc, "Corresponding author: Prof. Leo Joskowicz — josko@cs.huji.ac.il", size=9)
    legend_box(doc)

    doc.add_heading("Abstract", level=1)
    labeled(doc, "Purpose:",
            "Comparison of a current chest radiograph with a prior study to identify interval change is one of the most "
            "frequent tasks in thoracic imaging, yet the reliability of this judgment across readers is poorly quantified. "
            "We aimed to quantify the inter-observer agreement and variability of radiologists in detecting and "
            "characterizing interval change on longitudinal chest radiograph (CXR) pairs, and to determine how many "
            "reader subgroup size is associated with coverage of the full-panel reference.")
    labeled(doc, "Materials and Methods:",
            "Five [DECIDE|board-certified] radiologists independently reviewed 100 longitudinal CXR pairs (a prior and a "
            "current image from the same patient) [DATA|ICU/inpatient cohort; retrospective; IRB approval and waiver of "
            "consent obtained]. Readers marked every region of interval change and classified it as appearance (new), "
            "disappearance (resolved), or persistence with a change in size and/or intensity; findings were grouped into "
            "positive (new/worsening) and negative (resolved/improving) change. A region marked by at least k readers was "
            "assigned consensus level k. Agreement was quantified with a Pairwise Agreement Index (PAI) at the detection "
            "and pair levels and with per-reader leave-one-out sensitivity by consensus level. To determine the number of "
            "readers required, we measured, for every subgroup of 1–5 readers, the fraction of the full five-reader "
            "reference findings the subgroup recovered (coverage), averaged over all subgroups of each size.")
    labeled(doc, "Results:",
            "Readers marked a mean of 1.68 change findings per pair (range 1.45–1.97). The mean pairwise agreement between "
            "two readers was 0.55 at the detection level (range 0.46–0.64) and 0.62 at the pair level (range 0.52–0.70); "
            "agreement was higher for positive than for negative change. Of 185 positive change findings, 45% were "
            "reported by a single reader and only 15% by all five (negative: 40% solo, 9% unanimous). Per-reader "
            "leave-one-out sensitivity rose with consensus, from 0.46 at level 1 to 0.89 at level 4 for positive change. A "
            "single reader recovered a mean of 49% of the full five-reader reference; three readers recovered 81% "
            "(range across triplets, 71%–91%).")
    labeled(doc, "Conclusions:",
            "Radiologist agreement on interval change on chest radiographs is moderate and many findings are reported by "
            "only one reader. Three readers provided a pragmatic trade-off between annotation effort and coverage, "
            "recovering a mean of 81% of findings identified by the full five-reader panel.")

    p = doc.add_paragraph()
    p.add_run("Key Points").bold = True
    bullets(doc, [
        "Inter-observer agreement for interval change on chest radiographs is moderate (mean pairwise detection agreement "
        "0.55) and highest for new/worsening findings.",
        "Nearly half of all change findings (45% positive, 40% negative) were reported by only one of five readers.",
        "Three-reader subgroups recovered a mean of 81% of the five-reader reference (range, 71%–91%); full coverage "
        "was observed only for the five-reader panel.",
    ])

    doc.add_heading("1. Introduction", level=1)
    body(doc,
         "Serial comparison of chest radiographs is among the highest-volume interpretive tasks in radiology, "
         "particularly in intensive-care and inpatient settings, where daily films are used to track lines and tubes, "
         "effusions, consolidation, edema, and pneumothorax. Clinical decisions frequently turn on whether a finding is "
         "new, resolved, larger, or smaller relative to the prior study.")
    body(doc,
         "Despite the ubiquity of this task, the reliability of interval-change judgments across radiologists has "
         "received far less attention than single-image detection or quantitative measurement. Prior inter-observer "
         "studies in CT have shown that variability is large and that two or three observers may not capture its full "
         "range [DECIDE|Joskowicz 2019 Eur Radiol; Olesinski 2026 J Thorac Imaging]. Whether the same holds for the "
         "comparative, detection-level task of interval change on radiographs is unknown, yet it directly determines what "
         "should count as ground truth and how many readers are needed to define it.")
    body(doc,
         "In this study we quantify the inter-observer agreement and variability of five radiologists detecting interval "
         "change on 100 longitudinal CXR pairs, and we introduce a subgroup analysis that measures how completely groups "
         "of one to five readers reproduce the full-panel reference.")

    doc.add_heading("2. Materials and Methods", level=1)
    doc.add_heading("2.1 Study population", level=2)
    body(doc,
         "[DECIDE|Retrospective] analysis of 100 longitudinal chest radiograph pairs from [DATA|N patients] imaged at "
         "[DATA|institution] between [DATA|dates]. Each pair comprised a prior and a current [DECIDE|AP/portable] "
         "radiograph from the same patient [DATA|selection criteria: consecutive ICU studies with an interpretable prior "
         "within X days]. [DATA|IRB approval number; informed consent waived.] Patient demographics are summarized in "
         "[SUGGEST|Table S1].")
    doc.add_heading("2.2 Readers and reading protocol", level=2)
    body(doc,
         "Five radiologists ([DECIDE|R1–R5]; [DATA|years of experience]) independently reviewed all 100 pairs on a "
         "dedicated annotation workstation, blinded to one another and to clinical information beyond the two images. For "
         "each pair, a reader delineated every region of interval change with an elliptical marker and labeled it "
         "Appearance (new), Disappearance (resolved), or Persistence (present on both, with an annotated change in size "
         "and intensity). Each region was reduced to a signed change: positive (appearance, or persistence with increased "
         "size/intensity) or negative (disappearance, or persistence with decreased size/intensity). Persistence without "
         "change was excluded.")
    doc.add_heading("2.3 Reference standard and consensus levels", level=2)
    body(doc,
         "Annotations were rasterized to the current-image space and grouped into connected components (findings). Each "
         "finding was assigned a consensus level equal to the number of readers (1–5) who independently marked an "
         "overlapping region. Consensus level 1 is the union of all reader findings; consensus level 5 comprises only "
         "unanimous findings.")
    doc.add_heading("2.4 Agreement, variability, and subgroup metrics", level=2)
    bullets(doc, [
        "Pairwise Agreement Index (PAI): for two readers, the per-detection PAI is 2M/(2M + U), where M is the number of "
        "matched (overlapping) findings and U the number of unmatched findings; the per-pair PAI averages agreement over "
        "the 100 image pairs.",
        "Human-Matched Detection Rate (HMDR): the fraction of a reader's findings overlapping a finding of at least one "
        "other reader; Unmatched Detections Per Pair (UDPP): the mean number of a reader's solo findings per pair.",
        "Sensitivity at consensus level (leave-one-out): for each reader, the fraction of the other four readers' "
        "consensus findings, at each consensus level, that the reader also marked.",
        "Reader-subgroup coverage: for every subgroup of size m (1–5), the fraction of the full five-reader reference "
        "findings recovered by the subgroup (a finding is recovered if any subgroup member marked it), averaged over all "
        "C(5, m) subgroups. Coverage is 1.0 by construction at m = 5; its rise with m quantifies how many readers are "
        "needed to characterize the coverage-effort relationship.",
    ])
    doc.add_heading("2.5 Statistical analysis", level=2)
    body(doc,
         "Agreement metrics are reported as means with ranges across reader pairs; subgroup coverage as the mean with the "
         "min–max envelope across subgroups of each size. [SUGGEST|Bootstrap 95% confidence intervals over image pairs to "
         "be added.] Analyses were performed in Python [DATA|3.x] with NumPy, SciPy, and scikit-image.")

    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1 Finding burden", level=2)
    body(doc,
         "Across 100 pairs the five readers marked a mean of 1.68 change findings per pair, with per-reader totals of "
         "145–197 (mean 1.45–1.97 per pair; maximum 4–6 on a single pair)—a 1.36-fold difference in reporting rate "
         "between the least and most prolific reader (Table 1).")
    body(doc, "Table 1. Change-finding burden per reader (100 pairs).")
    table(doc, ["Reader", "Total findings", "Mean per pair", "SD", "Max per pair"], [
        ["R1", 148, 1.48, 0.77, 4], ["R2", 185, 1.85, 1.04, 6], ["R3", 145, 1.45, 0.80, 4],
        ["R4", 166, 1.66, 0.95, 5], ["R5", 197, 1.97, 0.84, 4],
    ])

    doc.add_heading("3.2 Pairwise agreement", level=2)
    body(doc,
         "The mean per-detection PAI between two readers was 0.55 (range 0.46–0.64); the mean per-pair PAI was 0.62 "
         "(range 0.52–0.70). The highest-agreeing reader pair concurred on only ~64% of individual detections, and the "
         "lowest on ~46% (Table 2, Figure 1). Agreement was higher for positive change (mean per-detection PAI 0.57) than "
         "for negative change (0.51), indicating that new or worsening findings are marked more reproducibly than "
         "resolving ones.")
    body(doc, "Table 2. Pairwise Agreement Index (per detection, all change).")
    table(doc, ["", "R1", "R2", "R3", "R4", "R5"], [
        ["R1", "—", 0.47, 0.59, 0.51, 0.59], ["R2", 0.47, "—", 0.50, 0.46, 0.50],
        ["R3", 0.59, 0.50, "—", 0.57, 0.63], ["R4", 0.51, 0.46, 0.57, "—", 0.64],
        ["R5", 0.59, 0.50, 0.63, 0.64, "—"],
    ])
    figure(doc, HUM / "per_label_agreement_all_humans.png",
           "Figure 1a. Pairwise Agreement Index per detection (all change), five radiologists.", 4.3)
    figure(doc, HUM / "per_pair_agreement_all_humans.png",
           "Figure 1b. Pairwise Agreement Index per pair (all change), five radiologists.", 4.3)

    doc.add_heading("3.3 Consensus structure and disagreement", level=2)
    body(doc,
         "Disagreement was dominated by low-consensus findings. Of 185 positive change findings, 101 (55%) reached "
         "consensus level 2, 76 (41%) level 3, 52 (28%) level 4, and only 27 (15%) were unanimous; 45% (83/185) were solo "
         "findings marked by a single reader. Negative change followed the same pattern: of 139 findings, 84 (60%) reached "
         "level 2 and only 13 (9%) were unanimous, with 40% (55/139) solo (Figure 2). Between 3% and 28% of each reader's "
         "findings were unmatched by any peer (HMDR 0.72–0.97; Table 4).")
    figure(doc, HUM / "recall_by_consensus_level_humans.png",
           "Figure 2. Per-reader leave-one-out sensitivity as a function of consensus level, for positive (left) and "
           "negative (right) change. Sensitivity rises steeply as more readers agree.", 6.5)
    body(doc, "Table 4. Corroboration of each reader's findings.")
    table(doc, ["Reader", "HMDR (pos)", "HMDR (neg)", "UDPP (pos)", "UDPP (neg)"], [
        ["R1", 0.83, 0.85, 0.15, 0.09], ["R2", 0.72, 0.78, 0.30, 0.17], ["R3", 0.97, 0.88, 0.02, 0.08],
        ["R4", 0.89, 0.88, 0.11, 0.08], ["R5", 0.77, 0.83, 0.26, 0.15],
    ])

    doc.add_heading("3.4 Agreement rises with consensus", level=2)
    body(doc,
         "Per-reader leave-one-out sensitivity increased monotonically with consensus level (Figure 2, Table 3). Averaged "
         "across readers, sensitivity for positive change rose from 0.46 at consensus level 1 to 0.72, 0.80, and 0.89 at "
         "levels 2–4; for negative change from 0.46 to 0.62, 0.82, and 0.85. When only one other reader had seen a "
         "finding, a given radiologist reproduced it less than half the time; when all four others agreed, ~85–90% of the "
         "time.")
    body(doc, "Table 3. Mean per-reader leave-one-out sensitivity by consensus level.")
    table(doc, ["Consensus level", "Positive", "Negative"], [
        ["≥1", 0.46, 0.46], ["≥2", 0.72, 0.62], ["≥3", 0.80, 0.82], ["≥4", 0.89, 0.85],
    ])

    doc.add_heading("3.5 Reader-subgroup coverage", level=2)
    c_all = {int(m): cov[m]["all"] for m in cov}
    body(doc,
         f"A single radiologist recovered a mean of {c_all[1]['mean']*100:.0f}% "
         f"(range {c_all[1]['min']*100:.0f}–{c_all[1]['max']*100:.0f}%) of the findings identified by the full "
         f"five-reader panel. Coverage rose to {c_all[2]['mean']*100:.0f}% with two readers, "
         f"{c_all[3]['mean']*100:.0f}% with three, and {c_all[4]['mean']*100:.0f}% with four (Figure 3, Table 5). The "
         f"marginal gain from each added reader was +{mg['2']*100:.0f}% for the second reader, "
         f"+{mg['3']*100:.0f}% for the third, +{mg['4']*100:.0f}% for the fourth, and +{mg['5']*100:.0f}% for the fifth. "
         f"Three-reader subgroups ranged from {c_all[3]['min']*100:.0f}% to {c_all[3]['max']*100:.0f}% coverage. "
         "Positive and negative change showed similar coverage patterns (Figure 3).")
    body(doc, "Table 5. Reader-subgroup coverage of the full five-reader reference (all change).")
    table(doc, ["# readers", "Coverage (mean)", "Range (min–max)", "Marginal gain"], [
        ["1", f"{c_all[1]['mean']:.2f}", f"{c_all[1]['min']:.2f}–{c_all[1]['max']:.2f}", "—"],
        ["2", f"{c_all[2]['mean']:.2f}", f"{c_all[2]['min']:.2f}–{c_all[2]['max']:.2f}", f"+{mg['2']*100:.0f}%"],
        ["3", f"{c_all[3]['mean']:.2f}", f"{c_all[3]['min']:.2f}–{c_all[3]['max']:.2f}", f"+{mg['3']*100:.0f}%"],
        ["4", f"{c_all[4]['mean']:.2f}", f"{c_all[4]['min']:.2f}–{c_all[4]['max']:.2f}", f"+{mg['4']*100:.0f}%"],
        ["5", f"{c_all[5]['mean']:.2f}", "—", f"+{mg['5']*100:.0f}%"],
    ])
    figure(doc, HUM / "reader_subgroup_coverage.png",
           "Figure 3. Coverage of the full five-reader reference as a function of subgroup size, for all, positive, and "
           "negative change. Points are means over all subgroups of each size; shaded bands show the min–max envelope. "
           "Coverage rises with diminishing returns and, by construction, reaches 1.0 only with all five readers.", 6.2)

    # 3.6 subgroup HMDR / UDPP
    doc.add_heading("3.6 Corroboration and solo findings by subgroup size", level=2)
    hmdr = {int(m): metrics["subgroup_hmdr"][m]["all"] for m in metrics["subgroup_hmdr"]}
    udpp = {int(m): metrics["subgroup_udpp"][m]["all"] for m in metrics["subgroup_udpp"]}
    body(doc,
         f"We next asked how corroborated a subgroup's reads are as readers are added. Within a subgroup, the "
         f"Human-Matched Detection Rate (HMDR)—the fraction of a member's findings independently marked by at least one "
         f"other member—rose from {hmdr[2]['mean']:.2f} for reader pairs to {hmdr[3]['mean']:.2f} for triplets, "
         f"{hmdr[4]['mean']:.2f} for quadruplets, and {hmdr[5]['mean']:.2f} for the full panel (a lone reader has no "
         f"corroboration by definition). Conversely, the number of unmatched (solo) findings per reader per pair (UDPP) "
         f"fell from {udpp[1]['mean']:.2f} for a single reader to {udpp[2]['mean']:.2f} for pairs, "
         f"{udpp[3]['mean']:.2f} for triplets, and {udpp[5]['mean']:.2f} for the full panel (Figure 4, Table 6).")
    body(doc, "Table 6. Subgroup corroboration (HMDR) and solo-finding rate (UDPP), all change (mean [min–max]).")
    table(doc, ["# readers", "HMDR (mean)", "HMDR range", "UDPP (mean)", "UDPP range"], [
        [str(m),
         f"{hmdr[m]['mean']:.2f}", f"{hmdr[m]['min']:.2f}–{hmdr[m]['max']:.2f}",
         f"{udpp[m]['mean']:.2f}", f"{udpp[m]['min']:.2f}–{udpp[m]['max']:.2f}"]
        for m in range(1, 6)
    ])
    figure(doc, HUM / "reader_subgroup_hmdr_udpp.png",
           "Figure 4. Within-subgroup corroboration (HMDR, left) and solo-finding rate (UDPP, right) as a function of "
           "subgroup size. As readers are added, findings are increasingly corroborated and solo findings decline, with "
           "the largest change occurring up to three readers.", 6.5)

    doc.add_heading("4. Discussion", level=1)
    body(doc,
         "In a multi-reader study of interval change on 100 longitudinal chest radiographs, radiologist agreement was "
         "moderate, with many low-consensus findings: nearly half of all changes were reported by a single reader and fewer "
         "than one in six were unanimous. Sensitivity increased with the number of independent readers identifying a finding.")
    body(doc,
         "Reader-subgroup coverage increased from 68.6% for two-reader subsets to 81.3% for three-reader subsets; adding "
         "a fourth reader increased mean coverage by 10.2 percentage points. The 71%–91% coverage range across triplets "
         "shows that performance depends on which readers are selected. No threshold for an adequate reference standard "
         "was predefined, and full coverage is expected only when all five readers define the reference. These data support "
         "three readers as a pragmatic compromise for this cohort, rather than demonstrating an adequate or universally "
         "sufficient panel size.")
    body(doc,
         "For quality assurance, the high proportion of solo findings indicates that the reported pattern of interval change "
         "can differ materially between readers. For reference-standard construction, a single-reader annotation does not "
         "represent the full five-reader union in this cohort. Reporting consensus level and the panel composition makes "
         "that uncertainty visible. [SUGGEST|A companion analysis evaluates an automated change-detection method against "
         "this consensus.]")
    body(doc,
         "Limitations. This was a [DECIDE|single-center, retrospective] study of 100 pairs read by five radiologists. "
         "Reader experience was [DATA|not stratified]; annotations used elliptical region markers rather than pixel-wise "
         "contours, which may merge or split adjacent findings. The cohort was enriched for [DECIDE|ICU studies], which "
         "may over-represent devices and rapidly evolving findings and limit generalizability. Finally, the reference is "
         "defined by the readers themselves; consensus level and subgroup coverage are pragmatic surrogates for truth, "
         "not an independent gold standard, and coverage necessarily reaches 1.0 only at the full panel size.")

    doc.add_heading("5. Conclusions", level=1)
    body(doc,
         "Radiologist agreement on interval change on chest radiographs is moderate and graded, and nearly half of all "
         "identified changes are seen by a single reader. In this cohort, three readers recovered a mean of 81% of the "
         "five-reader reference, representing a pragmatic trade-off between annotation effort and coverage rather than a "
         "formally established threshold for a reference standard.")

    doc.add_heading("Supplementary figures", level=1)
    body(doc, "Per-reader sensitivity-at-consensus curves and the positive/negative PAI heatmaps:")
    for fn, cap in [
        ("per_label_agreement_pos_humans.png", "Supplementary Figure S1. PAI per detection (positive change)."),
        ("per_label_agreement_neg_humans.png", "Supplementary Figure S2. PAI per detection (negative change)."),
        ("per_pair_agreement_pos_humans.png", "Supplementary Figure S3. PAI per pair (positive change)."),
        ("per_pair_agreement_neg_humans.png", "Supplementary Figure S4. PAI per pair (negative change)."),
    ]:
        figure(doc, HUM / fn, cap, 3.6)

    center(doc, "Draft v3 (human readers only). Reader labels R1–R5 are anonymized. "
                "Source metrics: ov_results_human_only/human_ov_metrics.json.", italic=True, size=8)

    try:
        doc.save(str(OUT_DOCX))
        print(f"Wrote {OUT_DOCX}")
    except PermissionError:
        alt = OUT_DOCX.with_name(OUT_DOCX.stem + "_rebuilt.docx")
        doc.save(str(alt))
        print(f"WARNING: {OUT_DOCX.name} is open/locked; wrote {alt} instead. "
              f"Close the open file and re-run to overwrite the original.")


if __name__ == "__main__":
    build()
