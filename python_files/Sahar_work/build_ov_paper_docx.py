"""Build the Word (.docx) version of the observer-variability paper draft, with
embedded figures from the OV results. Also generates the consensus-level
distribution figure (Fig 4) from the consensus counts.

Run (from repo root):
    python python_files/Sahar_work/build_ov_paper_docx.py
"""
from __future__ import annotations

from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
import numpy as np  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Pt, Inches, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402


REPO = Path(__file__).resolve().parents[1]
SQ = REPO / "Sahar_work" / "files" / "ov_results_sq" / "no_cc_itamar_plus_lmm5_100"
LOO = REPO / "Sahar_work" / "files" / "ov_results_main_loo_itamar_plus_lmm5"
OUT_DOCX = REPO / "Sahar_work" / "OV_paper_draft_JThoracImaging.docx"


# ── Figure 4: consensus-level distribution (fraction of findings by # readers) ──
def make_consensus_distribution(out_path: Path) -> None:
    pos_ge = np.array([185, 101, 76, 52, 27], dtype=float)
    neg_ge = np.array([139, 84, 49, 32, 13], dtype=float)
    pos_exact = np.array([pos_ge[k] - (pos_ge[k + 1] if k + 1 < 5 else 0) for k in range(5)])
    neg_exact = np.array([neg_ge[k] - (neg_ge[k + 1] if k + 1 < 5 else 0) for k in range(5)])
    pos_frac = pos_exact / pos_ge[0]
    neg_frac = neg_exact / neg_ge[0]

    x = np.arange(1, 6)
    w = 0.38
    fig, ax = plt.subplots(figsize=(8, 5), dpi=180)
    ax.bar(x - w / 2, pos_frac, w, label="Positive change", color="#003366")
    ax.bar(x + w / 2, neg_frac, w, label="Negative change", color="#800000")
    for xi, v in zip(x - w / 2, pos_frac):
        ax.text(xi, v + 0.005, f"{v*100:.0f}%", ha="center", va="bottom", fontsize=8)
    for xi, v in zip(x + w / 2, neg_frac):
        ax.text(xi, v + 0.005, f"{v*100:.0f}%", ha="center", va="bottom", fontsize=8)
    ax.set_xticks(x)
    ax.set_xlabel("Number of readers who marked the finding (consensus level)")
    ax.set_ylabel("Fraction of change findings")
    ax.set_title("Distribution of Change Findings by Consensus Level", fontweight="bold")
    ax.grid(True, axis="y", linestyle=":", alpha=0.6)
    ax.legend(frameon=True, edgecolor="black")
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)


# ── docx helpers ──────────────────────────────────────────────────────────────
def set_base_style(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)


def add_center(doc: Document, text: str, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)


def add_labeled(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label} ")
    r.bold = True
    p.add_run(body)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)


def add_figure(doc: Document, path: Path, caption: str, width_in: float) -> None:
    if not path.exists():
        doc.add_paragraph(f"[MISSING FIGURE: {path.name}]")
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)


# ── Build ────────────────────────────────────────────────────────────────────
def build() -> None:
    fig4 = LOO / "consensus_level_distribution.png"
    make_consensus_distribution(fig4)

    doc = Document()
    set_base_style(doc)

    add_title(doc, "Inter-observer Variability in the Detection of Interval Change on "
                   "Longitudinal Chest Radiographs: How Many Readers Define the Truth?")
    add_center(doc, "[Author list]", size=11)
    add_center(doc, "[School of Computer Science and Engineering, The Hebrew University of Jerusalem; "
                    "Department of Radiology, Hadassah Hebrew University Medical Center, Jerusalem, Israel]",
               italic=True, size=9)
    add_center(doc, "Corresponding author: Prof. Leo Joskowicz — josko@cs.huji.ac.il", size=9)

    doc.add_heading("Abstract", level=1)
    add_labeled(doc, "Purpose:",
                "Comparison of a current chest radiograph with a prior study to identify interval change is one "
                "of the most frequent tasks in thoracic imaging, yet the reliability of this judgment across readers "
                "is poorly quantified. We aimed to quantify the inter-observer agreement and variability of "
                "radiologists in detecting and characterizing interval change on longitudinal chest radiograph (CXR) "
                "pairs, and to determine how many readers are required to establish a stable reference standard.")
    add_labeled(doc, "Materials and Methods:",
                "Five [board-certified] radiologists independently reviewed 100 longitudinal CXR pairs (a prior and "
                "a current image from the same patient) [ICU/inpatient cohort; retrospective; IRB approval and waiver "
                "of consent obtained]. For each pair, readers marked every region of interval change and classified it "
                "as appearance (new finding), disappearance (resolved finding), or persistence with a change in size "
                "and/or intensity (increase/decrease). Findings were grouped into positive change (new or worsening) "
                "and negative change (resolved or improving). A region marked by at least k readers was assigned "
                "consensus level k. Agreement was quantified with a Pairwise Agreement Index (PAI) at the level of "
                "individual detections and whole pairs, the Human-Matched Detection Rate (HMDR), the number of "
                "Unmatched Detections Per Pair (UDPP), and the per-reader leave-one-out sensitivity as a function of "
                "consensus level.")
    add_labeled(doc, "Results:",
                "Readers marked a mean of 1.68 change findings per pair (range across readers, 1.45–1.97; 145–197 "
                "findings each). The mean pairwise agreement between two readers was 0.55 at the detection level "
                "(range 0.46–0.64) and 0.61 at the pair level (range 0.52–0.70); agreement was higher for positive "
                "than for negative change (0.57 vs 0.51). Of 185 distinct positive change findings, 45% (83/185) were "
                "reported by a single reader and only 15% (27/185) by all five; for negative change, 40% (55/139) were "
                "solo and 9% (13/139) unanimous. Per-reader leave-one-out sensitivity against the remaining panel rose "
                "monotonically with consensus level, from 0.46 at consensus level 1 to 0.89 at level 4 for positive "
                "change (0.46 to 0.85 for negative). Between 3% and 28% of each reader's findings were unmatched by any peer.")
    add_labeled(doc, "Conclusions:",
                "Radiologist agreement on interval change on chest radiographs is moderate and is dominated by "
                "low-consensus findings: nearly half of all identified changes were seen by only one reader. Agreement "
                "rises steeply with consensus, so two or three readers capture only a fraction of the clinically "
                "reported change. A reference standard for interval change—and for the evaluation of automated "
                "change-detection algorithms—should be built from a multi-reader consensus rather than from any single reader.")

    p = doc.add_paragraph()
    p.add_run("Key Points").bold = True
    add_bullets(doc, [
        "Inter-observer agreement for interval change on chest radiographs is moderate (mean pairwise detection "
        "agreement 0.55) and highest for new/worsening findings.",
        "Almost half of all change findings (45% positive, 40% negative) were reported by only one of five readers; "
        "fewer than one in six were unanimous.",
        "Reader sensitivity against the panel nearly doubles from single-reader to four-reader consensus, indicating "
        "that two or three readers are insufficient to define the full range of reported change.",
    ])

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    add_body(doc,
             "Serial comparison of chest radiographs is among the highest-volume interpretive tasks in radiology, "
             "particularly in the intensive-care and inpatient settings, where daily films are used to track lines and "
             "tubes, effusions, consolidation, edema, and pneumothorax. Clinical decisions—escalation of care, drainage, "
             "antibiotic changes, device repositioning—frequently turn on whether a finding is new, resolved, larger, or "
             "smaller relative to the prior study.")
    add_body(doc,
             "Despite the ubiquity of this task, the reliability of interval-change judgments across radiologists has "
             "received far less attention than the reliability of single-image detection or of quantitative measurement. "
             "Prior work from our group and others has shown that inter-observer variability in manual delineation and "
             "measurement is large and that two or three observers may not be sufficient to capture its full range "
             "[Joskowicz 2019 Eur Radiol; Olesinski 2026 J Thorac Imaging]. Whether the same holds for the fundamentally "
             "comparative, detection-level task of interval change on radiographs is unknown, yet it directly determines "
             "what should count as ground truth when training and evaluating automated longitudinal analysis tools.")
    add_body(doc,
             "In this study we quantify the inter-observer agreement and variability of five radiologists detecting and "
             "characterizing interval change on 100 longitudinal CXR pairs. We introduce a consensus-level framework in "
             "which every change region is labeled by the number of readers who identified it, and we use group-wise "
             "agreement metrics to answer a practical question: how many readers are needed before the reference standard "
             "stabilizes?")

    # 2. Methods
    doc.add_heading("2. Materials and Methods", level=1)
    doc.add_heading("2.1 Study population", level=2)
    add_body(doc,
             "[Retrospective] analysis of 100 longitudinal chest radiograph pairs from [N patients] imaged at "
             "[institution] between [dates]. Each pair comprised a prior and a current [AP/portable] radiograph from the "
             "same patient [selection criteria: consecutive ICU studies with an interpretable prior within X days]. "
             "[IRB approval number; informed consent waived.] Patient demographics are summarized in [Table S1].")
    doc.add_heading("2.2 Readers and reading protocol", level=2)
    add_body(doc,
             "Five radiologists ([R1–R5]; [years of experience: …]) independently reviewed all 100 pairs on a dedicated "
             "annotation workstation, blinded to one another and to any clinical information beyond the two images. For "
             "each pair, a reader delineated every region of interval change with an elliptical marker and assigned one "
             "of the following labels: Appearance (present on the current image but not the prior; new); Disappearance "
             "(present on the prior but not the current; resolved); Persistence (present on both, with an annotated change "
             "in size and intensity). For analysis, each marked region was reduced to a signed change: positive change "
             "(appearance, or persistence with increased size/intensity) or negative change (disappearance, or persistence "
             "with decreased size/intensity). Persistence without change was excluded.")
    doc.add_heading("2.3 Reference standard and consensus levels", level=2)
    add_body(doc,
             "Reader annotations were rasterized to the coordinate space of the current image and grouped into connected "
             "components (findings). For each finding we computed its consensus level—the number of readers (1–5) who "
             "independently marked an overlapping region. A finding at consensus level k was identified by at least k "
             "readers. This yields a graded reference standard: consensus level 1 is the union of all reader findings, and "
             "consensus level 5 comprises only unanimous findings.")
    doc.add_heading("2.4 Agreement and variability metrics", level=2)
    add_body(doc, "All metrics were computed separately for positive and negative change.")
    add_bullets(doc, [
        "Pairwise Agreement Index (PAI): for each pair of readers, the per-detection PAI is 2M/(2M + U), where M is the "
        "number of matched (spatially overlapping) findings and U the number of unmatched findings—an F1-/Dice-like "
        "agreement on detections. The per-pair PAI averages agreement over the 100 image pairs.",
        "Human-Matched Detection Rate (HMDR): the fraction of a reader's findings that overlap a finding of at least one "
        "other reader (a precision-like measure of how corroborated a reader's marks are).",
        "Unmatched Detections Per Pair (UDPP): the mean number of a reader's findings, per image pair, not matched by any "
        "other reader (solo findings).",
        "Sensitivity at consensus level (leave-one-out): for each reader, the fraction of the consensus findings of the "
        "other four readers, at each consensus level, that the reader also marked. Because the scored reader is excluded "
        "from the reference, this is an unbiased, symmetric measure of how well each reader recovers the panel consensus.",
        "Consensus-based specificity: among image pairs in which the remaining panel identified no change of a given sign, "
        "the fraction in which the reader likewise marked none.",
    ])
    doc.add_heading("2.5 Statistical analysis", level=2)
    add_body(doc,
             "Agreement metrics are reported as means with ranges across reader pairs. Sensitivity is reported as a "
             "function of consensus level with per-reader curves. [Bland–Altman / 95% CIs / bootstrap resampling over "
             "pairs to be added.] Analyses were performed in Python [3.x] with NumPy, SciPy, and scikit-image.")

    # 3. Results
    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1 Finding burden", level=2)
    add_body(doc,
             "Across 100 pairs the five readers marked a mean of 1.68 change findings per pair. Per-reader totals ranged "
             "from 145 to 197 findings (mean 1.45–1.97 per pair; maximum 4–6 findings on a single pair), a 1.36-fold "
             "difference in reporting rate between the least and most prolific reader (Table 1).")
    add_body(doc, "Table 1. Change-finding burden per reader (100 pairs).")
    add_table(doc, ["Reader", "Total findings", "Mean per pair", "SD", "Max per pair"], [
        ["R1", 148, 1.48, 0.77, 4],
        ["R2", 185, 1.85, 1.04, 6],
        ["R3", 145, 1.45, 0.80, 4],
        ["R4", 166, 1.66, 0.95, 5],
        ["R5", 197, 1.97, 0.84, 4],
    ])

    doc.add_heading("3.2 Pairwise agreement", level=2)
    add_body(doc,
             "The mean per-detection PAI between two readers was 0.55 (range 0.46–0.64); the mean per-pair PAI was 0.61 "
             "(range 0.52–0.70). The highest-agreeing reader pair still concurred on only ~64% of individual detections, "
             "and the lowest on ~46% (Table 2, Figure 1). Agreement was consistently higher for positive change "
             "(mean per-detection PAI 0.57) than for negative change (0.51), indicating that new or worsening findings are "
             "marked more reproducibly than resolving ones.")
    add_body(doc, "Table 2. Pairwise Agreement Index (per detection, positive + negative combined).")
    add_table(doc, ["", "R1", "R2", "R3", "R4", "R5"], [
        ["R1", "—", 0.47, 0.59, 0.51, 0.59],
        ["R2", 0.47, "—", 0.50, 0.46, 0.50],
        ["R3", 0.59, 0.50, "—", 0.57, 0.63],
        ["R4", 0.51, 0.46, 0.57, "—", 0.64],
        ["R5", 0.59, 0.50, 0.63, 0.64, "—"],
    ])
    add_figure(doc, SQ / "per_label_agreement_all.png",
               "Figure 1a. Pairwise Agreement Index per detection (all change). The lower-right cell (M) is an "
               "automated method shown for reference; H–H = human–human, M–H = model–human.", 4.3)
    add_figure(doc, SQ / "per_pair_agreement_all.png",
               "Figure 1b. Pairwise Agreement Index per pair (all change).", 4.3)

    doc.add_heading("3.3 Consensus structure and disagreement", level=2)
    add_body(doc,
             "Disagreement was dominated by low-consensus findings. Of 185 distinct positive change findings (consensus "
             "level ≥1), 101 (55%) reached consensus level 2, 76 (41%) level 3, 52 (28%) level 4, and only 27 (15%) were "
             "unanimous; conversely, 45% (83/185) were solo findings marked by a single reader. Negative change followed "
             "the same pattern: of 139 findings, 84 (60%) reached level 2 and only 13 (9%) were unanimous, with 40% "
             "(55/139) solo (Figure 2).")
    add_body(doc,
             "Consistent with this, between 3% and 28% of each reader's findings were unmatched by any peer (HMDR "
             "0.72–0.97 for positive change), and readers contributed a mean of 0.02–0.30 solo findings per pair (UDPP; "
             "Table 4). Two readers accounted for most solo findings, while one reader's marks were almost always "
             "corroborated (HMDR 0.97).")
    add_figure(doc, LOO / "consensus_level_distribution.png",
               "Figure 2. Distribution of change findings by consensus level (fraction of findings marked by exactly "
               "1–5 readers). Nearly half of all findings were solo.", 5.6)
    add_body(doc, "Table 4. Corroboration of each reader's findings.")
    add_table(doc, ["Reader", "HMDR (pos)", "HMDR (neg)", "UDPP (pos)", "UDPP (neg)"], [
        ["R1", 0.83, 0.85, 0.15, 0.09],
        ["R2", 0.72, 0.78, 0.30, 0.17],
        ["R3", 0.97, 0.88, 0.02, 0.08],
        ["R4", 0.89, 0.88, 0.11, 0.08],
        ["R5", 0.77, 0.83, 0.26, 0.15],
    ])

    doc.add_heading("3.4 Agreement rises with consensus", level=2)
    add_body(doc,
             "Per-reader leave-one-out sensitivity against the remaining panel increased monotonically with consensus "
             "level (Figure 3, Table 3). Averaged across readers, sensitivity for positive change rose from 0.46 at "
             "consensus level 1 to 0.72 (level 2), 0.80 (level 3), and 0.89 (level 4); for negative change it rose from "
             "0.46 to 0.62, 0.82, and 0.85. In other words, when only one other reader had seen a finding, a given "
             "radiologist reproduced it less than half the time; when all four others agreed, they reproduced it in "
             "~85–90% of cases.")
    add_body(doc, "Table 3. Mean per-reader leave-one-out sensitivity by consensus level.")
    add_table(doc, ["Consensus level", "Positive", "Negative"], [
        ["≥1", 0.46, 0.46],
        ["≥2", 0.72, 0.62],
        ["≥3", 0.80, 0.82],
        ["≥4", 0.89, 0.85],
    ])
    add_figure(doc, LOO / "recall_by_consensus_level.png",
               "Figure 3. Per-reader leave-one-out sensitivity as a function of consensus level, for positive (left) "
               "and negative (right) change. Solid black = automated method (matched leave-one-out); dashed grey = "
               "automated method vs the full five-reader panel.", 6.5)
    add_body(doc,
             "Consensus-based specificity was high for most readers (several readers marked no change on essentially all "
             "pairs the panel deemed unchanged), although the number of change-free pairs was small and this estimate "
             "should be interpreted with caution.")

    # 4. Discussion
    doc.add_heading("4. Discussion", level=1)
    add_body(doc,
             "In a multi-reader study of interval change on 100 longitudinal chest radiographs, we found that radiologist "
             "agreement is moderate and is dominated by low-consensus findings. Nearly half of all change findings were "
             "reported by a single reader, and fewer than one in six were unanimous. Agreement climbed steeply with "
             "consensus level—reader sensitivity against the panel nearly doubled from single-reader to four-reader "
             "consensus—demonstrating that the truth of interval change is graded rather than binary.")
    add_body(doc,
             "These results extend to the comparative, detection-level task of interval change the central message of "
             "prior inter-observer studies in CT: that two or even three observers may not be sufficient to establish the "
             "full range of inter-observer variability [Joskowicz 2019]. They also parallel the recent finding that, for "
             "mediastinal lymph-node assessment, disagreements are systematically larger than agreements and warrant a "
             "rethinking of dichotomous thresholds [Olesinski 2026]. Here, the analog of that message is that a "
             "single-reader reference standard captures only a fraction of what a panel of radiologists collectively "
             "identifies as change.")
    add_body(doc,
             "The clinical implications are twofold. First, for reader practice and quality assurance, the predominance of "
             "solo findings suggests that a substantial share of reported interval change is not reproducibly perceived, "
             "and that consensus review (or over-reading) is likely to change the reported picture in a non-trivial "
             "fraction of cases. Second, for automated longitudinal analysis, our findings argue strongly against "
             "evaluating change-detection algorithms against any single reader: an algorithm judged against one "
             "radiologist would be penalized for missing that reader's solo findings and rewarded or penalized "
             "inconsistently across readers. A graded, consensus-level reference standard—reporting algorithm sensitivity "
             "separately at each consensus level—provides a fairer and more informative benchmark. In our data, an "
             "automated change-detection method reached a panel agreement (per-detection PAI 0.46; leave-one-out "
             "sensitivity within the human range at every consensus level) comparable to the lower end of inter-reader "
             "agreement; full model results are reported separately.")
    add_body(doc,
             "Limitations. This was a [single-center, retrospective] study of 100 pairs read by five radiologists; the "
             "number of change-free pairs limited the precision of specificity estimates. Reader experience was "
             "[not stratified]; annotations used elliptical region markers rather than pixel-wise contours, which may "
             "merge or split adjacent findings. The cohort was enriched for [ICU studies], which may over-represent "
             "devices and rapidly evolving findings and limit generalizability to outpatient chest radiography. Finally, "
             "the reference standard is itself defined by the readers; consensus level is a pragmatic surrogate for truth, "
             "not an independent gold standard.")

    # 5. Conclusions
    doc.add_heading("5. Conclusions", level=1)
    add_body(doc,
             "Radiologist agreement on interval change on chest radiographs is moderate and graded: nearly half of all "
             "identified changes were seen by only one of five readers, and agreement rose sharply with the number of "
             "concurring readers. Two or three readers are insufficient to define the full range of reported change. "
             "Reference standards for interval change—and the evaluation of automated change-detection tools—should be "
             "built from multi-reader consensus and reported as a function of consensus level.")

    doc.add_heading("Supplementary figures", level=1)
    add_body(doc, "Per-reader sensitivity-at-consensus curves (readers A–E) and the automated method (M_ICU):")
    for name, letter in [("avi", "A"), ("benny", "B"), ("sigal", "C"), ("smadar", "D"), ("nitzan", "E")]:
        add_figure(doc, SQ / f"sensitivity_consensus_levels_{name}.png",
                   f"Supplementary Figure S{letter}. Sensitivity at consensus levels — reader {letter}.", 3.6)
    add_figure(doc, SQ / "sensitivity_consensus_levels.png",
               "Supplementary Figure S-M. Sensitivity at consensus levels — automated method (five levels).", 3.6)

    add_body(doc, "")
    add_center(doc,
               "Draft prepared for internal review. Reader labels R1–R5 are anonymized. "
               "Source metrics: ov_results_main_loo_itamar_plus_lmm5 and ov_results_sq/no_cc_itamar_plus_lmm5_100.",
               italic=True, size=8)

    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    build()
